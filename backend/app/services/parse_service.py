from __future__ import annotations

import csv
import hashlib
import io
import logging
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

import fitz
import pandas as pd
from docx import Document as DocxDocument
from fastapi import HTTPException, status
from sqlalchemy.orm import Session

from backend.app.core.config import settings
from backend.app.core.runtime import ensure_runtime_dirs
from backend.app.db.models import Document, DocumentParseResult


logger = logging.getLogger(__name__)
PARSED_STATUSES = {"parsed", "failed"}
HASH_CHUNK_SIZE = 1024 * 1024


@dataclass(frozen=True)
class ParsedText:
    text: str
    parser_name: str
    parser_version: str
    truncated: bool = False


def _api_error(status_code: int, code: str, message: str) -> HTTPException:
    return HTTPException(
        status_code=status_code,
        detail={
            "success": False,
            "code": code,
            "message": message,
            "data": None,
        },
    )


def _bad_request(message: str) -> HTTPException:
    return _api_error(status.HTTP_400_BAD_REQUEST, "BAD_REQUEST", message)


def _not_found(message: str) -> HTTPException:
    return _api_error(status.HTTP_404_NOT_FOUND, "NOT_FOUND", message)


def _server_error(code: str, message: str) -> HTTPException:
    return _api_error(status.HTTP_500_INTERNAL_SERVER_ERROR, code, message)


def calculate_file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for chunk in iter(lambda: file.read(HASH_CHUNK_SIZE), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _safe_original_file_path(document: Document) -> Path:
    files_root = (settings.data_path / "files").resolve()
    target_path = (settings.data_path / document.relative_path).resolve()
    if not target_path.is_relative_to(files_root):
        logger.error(
            "Refusing to parse document file outside data/files: document_id=%s relative_path=%s resolved_path=%s",
            document.id,
            document.relative_path,
            target_path,
        )
        raise _server_error("UNSAFE_FILE_PATH", "Document file path is outside the files directory.")
    return target_path


def safe_parsed_file_path(parse_result: DocumentParseResult) -> Path | None:
    if not parse_result.parsed_relative_path:
        return None
    parsed_root = settings.parsed_path.resolve()
    target_path = (settings.data_path / parse_result.parsed_relative_path).resolve()
    if not target_path.is_relative_to(parsed_root):
        logger.error(
            "Refusing to access parsed file outside data/parsed: document_id=%s relative_path=%s resolved_path=%s",
            parse_result.document_id,
            parse_result.parsed_relative_path,
            target_path,
        )
        raise _server_error("UNSAFE_PARSED_FILE_PATH", "Parsed file path is outside the parsed directory.")
    return target_path


def _parsed_output_paths(document_id: str) -> tuple[str, Path, Path]:
    now = datetime.now(timezone.utc)
    relative_dir = f"parsed/{now:%Y/%m}"
    relative_path = f"{relative_dir}/{document_id}.txt"
    target_dir = settings.data_path / relative_dir
    target_path = target_dir / f"{document_id}.txt"
    temp_path = target_dir / f"{document_id}.txt.tmp-{uuid4().hex}"
    if not target_path.resolve().is_relative_to(settings.parsed_path.resolve()):
        logger.error("Computed parsed output path outside data/parsed: document_id=%s path=%s", document_id, target_path)
        raise _server_error("UNSAFE_PARSED_FILE_PATH", "Parsed output path is outside the parsed directory.")
    target_dir.mkdir(parents=True, exist_ok=True)
    return relative_path, target_path, temp_path


def _limit_text(text: str) -> tuple[str, bool]:
    if len(text) <= settings.max_parse_chars:
        return text, False
    return text[: settings.max_parse_chars], True


def _decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _parse_text(path: Path, parser_name: str) -> ParsedText:
    text, truncated = _limit_text(_decode_text_bytes(path.read_bytes()))
    return ParsedText(text=text, parser_name=parser_name, parser_version="1", truncated=truncated)


def _parse_pdf(path: Path) -> ParsedText:
    parts: list[str] = []
    truncated = False
    with fitz.open(path) as pdf:
        for page_index in range(min(pdf.page_count, settings.max_parse_pages)):
            parts.append(pdf.load_page(page_index).get_text("text"))
            current = "\n".join(parts)
            if len(current) > settings.max_parse_chars:
                current, truncated = _limit_text(current)
                return ParsedText(text=current, parser_name="pymupdf", parser_version=fitz.VersionBind, truncated=truncated)
    text, limit_truncated = _limit_text("\n".join(parts))
    return ParsedText(text=text, parser_name="pymupdf", parser_version=fitz.VersionBind, truncated=truncated or limit_truncated)


def _parse_docx(path: Path) -> ParsedText:
    document = DocxDocument(path)
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    text, truncated = _limit_text("\n".join(parts))
    return ParsedText(text=text, parser_name="python-docx", parser_version="1.2", truncated=truncated)


def _parse_csv(path: Path) -> ParsedText:
    data = path.read_bytes()
    text_source = _decode_text_bytes(data)
    reader = csv.reader(io.StringIO(text_source))
    rows: list[str] = []
    for index, row in enumerate(reader):
        if index >= settings.max_parse_rows:
            break
        rows.append("\t".join(row))
    text, truncated = _limit_text("\n".join(rows))
    return ParsedText(text=text, parser_name="csv", parser_version="1", truncated=truncated)


def _parse_excel(path: Path, extension: str) -> ParsedText:
    engine = "openpyxl" if extension == ".xlsx" else "xlrd"
    sheet_texts: list[str] = []
    with pd.ExcelFile(path, engine=engine) as excel:
        for sheet_name in excel.sheet_names[: settings.max_parse_sheets]:
            frame = pd.read_excel(excel, sheet_name=sheet_name, nrows=settings.max_parse_rows, dtype=str)
            frame = frame.fillna("")
            sheet_texts.append(f"# Sheet: {sheet_name}")
            sheet_texts.append(frame.to_csv(sep="\t", index=False))
            current = "\n".join(sheet_texts)
            if len(current) > settings.max_parse_chars:
                text, truncated = _limit_text(current)
                return ParsedText(text=text, parser_name=f"pandas-{engine}", parser_version=pd.__version__, truncated=truncated)
    text, truncated = _limit_text("\n".join(sheet_texts))
    return ParsedText(text=text, parser_name=f"pandas-{engine}", parser_version=pd.__version__, truncated=truncated)


def parse_file(path: Path, extension: str) -> ParsedText:
    if extension in {".txt", ".md"}:
        return _parse_text(path, "plain-text" if extension == ".txt" else "markdown")
    if extension == ".pdf":
        return _parse_pdf(path)
    if extension == ".docx":
        return _parse_docx(path)
    if extension == ".csv":
        return _parse_csv(path)
    if extension in {".xlsx", ".xls"}:
        return _parse_excel(path, extension)
    raise _bad_request("Unsupported document extension for parsing.")


def _short_error(error: Exception) -> str:
    if isinstance(error, HTTPException):
        detail = error.detail
        if isinstance(detail, dict) and isinstance(detail.get("message"), str):
            return detail["message"][:500]
    return "Document parsing failed."


def _upsert_failed_result(db: Session, document: Document, source_sha256: str, error_message: str) -> DocumentParseResult:
    existing = db.get(DocumentParseResult, document.id)
    now = datetime.now(timezone.utc)
    if existing is None:
        existing = DocumentParseResult(document_id=document.id, status="failed", source_sha256=source_sha256)
        db.add(existing)
    existing.status = "failed"
    existing.source_sha256 = source_sha256
    existing.parsed_relative_path = None
    existing.parser_name = None
    existing.parser_version = None
    existing.content_sha256 = None
    existing.char_count = 0
    existing.truncated = False
    existing.error_message = error_message[:500]
    existing.parsed_at = now
    db.commit()
    db.refresh(existing)
    return existing


def parse_document(db: Session, document_id: str) -> DocumentParseResult:
    ensure_runtime_dirs()
    document = db.get(Document, document_id)
    if document is None:
        raise _not_found("Document not found.")
    if document.status != "uploaded":
        raise _bad_request("Archived documents cannot be parsed.")

    source_path = _safe_original_file_path(document)
    if not source_path.exists():
        error_message = "Original file is missing."
        result = _upsert_failed_result(db, document, document.sha256, error_message)
        raise _bad_request(error_message)

    current_sha256 = calculate_file_sha256(source_path)
    if current_sha256 != document.sha256:
        raise _bad_request("Original file checksum does not match the document record.")

    existing = db.get(DocumentParseResult, document.id)
    if existing is not None and existing.status == "parsed" and existing.source_sha256 == document.sha256:
        return existing

    relative_path, target_path, temp_path = _parsed_output_paths(document.id)
    wrote_final_file = False
    old_parsed_path = safe_parsed_file_path(existing) if existing is not None and existing.parsed_relative_path else None

    try:
        parsed = parse_file(source_path, document.extension)
        content_sha256 = hashlib.sha256(parsed.text.encode("utf-8")).hexdigest()
        temp_path.write_text(parsed.text, encoding="utf-8", newline="")
        temp_path.replace(target_path)
        wrote_final_file = True

        result = existing
        if result is None:
            result = DocumentParseResult(document_id=document.id, status="parsed", source_sha256=document.sha256)
            db.add(result)
        result.status = "parsed"
        result.source_sha256 = document.sha256
        result.parsed_relative_path = relative_path
        result.parser_name = parsed.parser_name
        result.parser_version = parsed.parser_version
        result.content_sha256 = content_sha256
        result.char_count = len(parsed.text)
        result.truncated = parsed.truncated
        result.error_message = None
        result.parsed_at = datetime.now(timezone.utc)
        db.commit()
        db.refresh(result)

        if old_parsed_path is not None and old_parsed_path != target_path:
            old_parsed_path.unlink(missing_ok=True)
        return result
    except HTTPException as error:
        db.rollback()
        temp_path.unlink(missing_ok=True)
        if wrote_final_file:
            target_path.unlink(missing_ok=True)
        logger.exception("Document parsing failed: document_id=%s", document.id)
        failed = _upsert_failed_result(db, document, document.sha256, _short_error(error))
        raise _bad_request(failed.error_message or "Document parsing failed.")
    except Exception as error:
        db.rollback()
        temp_path.unlink(missing_ok=True)
        if wrote_final_file:
            target_path.unlink(missing_ok=True)
        logger.exception("Document parsing failed: document_id=%s", document.id)
        failed = _upsert_failed_result(db, document, document.sha256, _short_error(error))
        raise _bad_request(failed.error_message or "Document parsing failed.")


def get_parse_result(db: Session, document_id: str) -> DocumentParseResult:
    document = db.get(Document, document_id)
    if document is None:
        raise _not_found("Document not found.")
    result = db.get(DocumentParseResult, document_id)
    if result is None:
        raise _not_found("Parse result not found.")
    return result
